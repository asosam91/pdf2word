#!/usr/bin/env python3
"""pdf_to_word_images.py — v1.4.1
====================================================================
✔ Corrige AttributeError: '_Header' object has no attribute 'clear_content'
   (python‑docx no expone ese método). Ahora se eliminan encabezados y pies
   borrando sus párrafos/tablas.
"""
from __future__ import annotations

import argparse, sys, logging
from io import BytesIO
from pathlib import Path
from typing import List

try:
    import tkinter as _tk
    from tkinter import filedialog as _fd
except ImportError:
    _tk = None; _fd = None

from pdf2docx import Converter  # type: ignore
import fitz                      # PyMuPDF
from docx import Document        # type: ignore
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.shape import WD_INLINE_SHAPE
from PIL import Image            # type: ignore
import pandas as pd              # type: ignore
import matplotlib.pyplot as plt  # type: ignore
plt.rcParams["figure.autolayout"] = True

# ────────────────────────────────────────────────────────────────────
# Helpers: GUI & logging
# ────────────────────────────────────────────────────────────────────

def _pick_pdf_gui() -> Path | None:
    if _tk is None:
        return None
    root = _tk.Tk(); root.withdraw()
    file = _fd.askopenfilename(title="Selecciona un PDF", filetypes=[("PDF", "*.pdf")])
    root.destroy(); return Path(file) if file else None


def _setup_logging(pdf_path: Path) -> None:
    log = pdf_path.with_name(f"{pdf_path.stem}_process.log")
    logging.basicConfig(level=logging.INFO,
                        format="%(asctime)s [%(levelname)s] %(message)s",
                        handlers=[logging.FileHandler(log, encoding="utf-8"), logging.StreamHandler()])
    logging.info("=== Inicio de proceso ===")

# ────────────────────────────────────────────────────────────────────
# PDF → Word & extracción de imágenes
# ────────────────────────────────────────────────────────────────────

def _pdf_to_word(pdf: Path) -> Path:
    out = pdf.with_suffix('.docx'); logging.info("Convirtiendo %s → %s", pdf.name, out.name)
    cv = Converter(str(pdf)); cv.convert(str(out)); cv.close(); return out


def _looks_like_chart(img: Image.Image) -> bool:
    thumb = img.resize((64, 64)).convert('RGB'); colors = thumb.getcolors(64*64) or []
    return len(colors) < 150


def _extract_images(pdf: Path, keep_all: bool) -> List[str]:
    doc, stem = fitz.open(pdf), pdf.stem; out: List[str] = []
    for p_idx, page in enumerate(doc, 1):
        for i_idx, (xref, *_r) in enumerate(page.get_images(full=True), 1):
            pix = fitz.Pixmap(doc, xref); pix = fitz.Pixmap(fitz.csRGB, pix) if pix.alpha else pix
            if not keep_all and not _looks_like_chart(Image.frombytes('RGB',[pix.width,pix.height],pix.samples)):
                continue
            label = 'img' if keep_all else 'chart'
            name = f"{stem}_p{p_idx}_{label}{i_idx}.png"; pix.save(name); out.append(name)
            logging.info("  • %s", name)
    logging.info("Total imágenes exportadas: %d", len(out)); return out

# ────────────────────────────────────────────────────────────────────
# DOCX post‑proceso: formato, tablas a PNG, filtra imágenes
# ────────────────────────────────────────────────────────────────────

def _clear_header_footer(part) -> None:
    """Borra párrafos y tablas de Header/Footer porque python‑docx no ofrece clear_content()."""
    for p in list(part.paragraphs):
        p._element.getparent().remove(p._element)
    for t in list(part.tables):
        t._element.getparent().remove(t._element)


def _df_to_png(df: pd.DataFrame, out: str) -> None:
    fig, ax = plt.subplots(figsize=(min(8, max(2, len(df.columns)*1.2)), 0.6+0.4*len(df)))
    ax.axis('off'); tbl = ax.table(cellText=df.values, colLabels=df.columns, loc='center', cellLoc='center')
    tbl.scale(1,1.2); fig.savefig(out, dpi=300, bbox_inches='tight'); plt.close(fig)


def _convert_tables(doc: Document, stem: str):
    if not doc.tables: return
    for idx, tbl in enumerate(list(doc.tables), 1):
        data = [[c.text.strip() for c in row.cells] for row in tbl.rows]
        if not data: continue
        out = f"{stem}_table{idx}.png"; _df_to_png(pd.DataFrame(data[1:], columns=data[0] if len(data)>1 else None), out)
        para = tbl._element.addnext(doc.add_paragraph()._p)  # placeholder
        doc.paragraphs[-1].add_run().add_picture(out, width=Inches(6))
        tbl._element.getparent().remove(tbl._element); logging.info("  • %s (tabla→imagen)", out)


def _filter_non_charts(doc: Document) -> int:
    rels, removed = doc.part._rels, 0
    for shp in list(doc.inline_shapes):
        if shp.type != WD_INLINE_SHAPE.PICTURE: continue
        img_bytes = rels[shp._inline.graphic.graphicData.pic.blipFill.blip.embed]._target._blob
        if _looks_like_chart(Image.open(BytesIO(img_bytes))): continue
        shp._inline.getparent().getparent().remove(shp._inline.getparent()); removed += 1
    return removed


def _postprocess(docx: Path, font: str, size: float, spacing: float, keep_charts_only: bool):
    doc = Document(str(docx)); stem = docx.stem
    for s in doc.sections: _clear_header_footer(s.header); _clear_header_footer(s.footer)
    normal = doc.styles['Normal']; normal.font.name = font; normal.font.size = Pt(size)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    for p in doc.paragraphs:
        p.paragraph_format.line_spacing = spacing
        for r in p.runs:
            r.font.name = font; r.font.size = Pt(size); r._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    _convert_tables(doc, stem)
    if keep_charts_only:
        removed = _filter_non_charts(doc); logging.info("Imágenes no gráficas eliminadas: %d", removed)
    doc.save(str(docx)); logging.info("Word formateado y guardado (%s)", docx.name)

# ────────────────────────────────────────────────────────────────────
# CLI & main
# ────────────────────────────────────────────────────────────────────

def main():
    ap = argparse.ArgumentParser(add_help=False)
    ap.add_argument('pdf', nargs='?', type=Path); ap.add_argument('--include-all-images', action='store_true')
    ap.add_argument('--font', default='Calibri'); ap.add_argument('--size', type=float, default=11)
    ap.add_argument('--spacing', type=float, default=1.0); ap.add_argument('-h','--help', action='help')
    a = ap.parse_args(); pdf = a.pdf or _pick_pdf_gui()
    if not pdf or not pdf.is_file() or pdf.suffix.lower()!='.pdf': print('[!] Selecciona un PDF válido'); sys.exit(1)
    pdf = pdf.resolve(); _setup_logging(pdf)
    try:
        docx = _pdf_to_word(pdf)
        _extract_images(pdf, a.include_all_images)
        _postprocess(docx, a.font, a.size, a.spacing, not a.include_all_images)
    except Exception as e:
        logging.exception('Error inesperado'); print('Ocurrió un error; revisa el log.'); sys.exit(2)
    print('✓ Proceso completado. Revisa el DOCX, PNGs y log.')

if __name__=='__main__':
    main()
