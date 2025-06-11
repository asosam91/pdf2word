#!/usr/bin/env python3
"""pdf_to_word_images.py
====================================================================
Conversor *todo‑en‑uno* (PDF ➜ Word + PNGs) — **v1.4**
Corrección de bug: el log ahora se crea correctamente sin error de sufijo.

Cambios rápidos v1.4
--------------------
✔️  Soluciona `ValueError: Invalid suffix '_process.log'` usando `with_name()`.
✔️  Si el nombre del PDF es `Informe.pdf`, el log será `Informe_process.log`.
"""
from __future__ import annotations

import argparse, sys, logging
from io import BytesIO
from pathlib import Path
from typing import List

try:
    import tkinter as _tk
    from tkinter import filedialog as _fd
except ImportError:
    _tk = None; _fd = None

from pdf2docx import Converter  # type: ignore
import fitz                      # PyMuPDF
from docx import Document        # type: ignore
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.shape import WD_INLINE_SHAPE
from PIL import Image            # type: ignore
import pandas as pd              # type: ignore
import matplotlib.pyplot as plt  # type: ignore
plt.rcParams["figure.autolayout"] = True

# ────────────────────────────────────────────────────────────────────
# Logging helpers
# ────────────────────────────────────────────────────────────────────

def _setup_logging(pdf_path: Path) -> None:
    log_file = pdf_path.with_name(f"{pdf_path.stem}_process.log")
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        handlers=[logging.FileHandler(log_file, encoding="utf-8"), logging.StreamHandler()],
    )
    logging.info("=== Inicio de proceso ===")

# ────────────────────────────────────────────────────────────────────
# GUI fallback
# ────────────────────────────────────────────────────────────────────

def _pick_pdf_gui() -> Path | None:
    if _tk is None:
        return None
    root = _tk.Tk(); root.withdraw()
    file = _fd.askopenfilename(title="Selecciona un PDF", filetypes=[("PDF", "*.pdf")])
    root.destroy()
    return Path(file) if file else None

# ────────────────────────────────────────────────────────────────────
# Conversion + extraction
# ────────────────────────────────────────────────────────────────────

def _pdf_to_word(pdf_path: Path) -> Path:
    word_path = pdf_path.with_suffix(".docx")
    logging.info(f"Convirtiendo {pdf_path.name} → {word_path.name}")
    cv = Converter(str(pdf_path)); cv.convert(str(word_path)); cv.close()
    return word_path


def _extract_charts(pdf_path: Path, keep_all: bool) -> List[str]:
    logging.info("Extrayendo imágenes (modo: %s)", "todas" if keep_all else "solo gráficas")
    doc = fitz.open(pdf_path); extracted: List[str] = []
    stem = pdf_path.stem
    for p_idx, page in enumerate(doc, 1):
        for img_idx, (xref, *_rest) in enumerate(page.get_images(full=True), 1):
            pix = fitz.Pixmap(doc, xref)
            if pix.alpha: pix = fitz.Pixmap(fitz.csRGB, pix)
            keep = True
            if not keep_all:
                keep = _looks_like_chart(Image.frombytes("RGB", [pix.width, pix.height], pix.samples))
            if not keep:
                continue
            label = "chart" if keep_all or _looks_like_chart(Image.frombytes("RGB", [pix.width, pix.height], pix.samples)) else "img"
            out_name = f"{stem}_p{p_idx}_{label}{img_idx}.png"
            pix.save(out_name); extracted.append(out_name)
            logging.info("  • %s", out_name)
    logging.info("%d imagen(es) exportadas", len(extracted))
    return extracted

# ────────────────────────────────────────────────────────────────────
# DOCX post‑processing: format + tables → images
# ────────────────────────────────────────────────────────────────────

def _postprocess_docx(
    docx_path: Path,
    font: str,
    size_pt: float,
    spacing: float,
    keep_only_charts: bool,
) -> None:
    logging.info("Aplicando formato y convirtiendo tablas …")
    doc = Document(str(docx_path))

    # headers/footers
    for sect in doc.sections:
        sect.header.clear_content(); sect.footer.clear_content()

    # base style
    normal = doc.styles["Normal"]; normal.font.name = font; normal.font.size = Pt(size_pt)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font)

    for p in doc.paragraphs:
        p.paragraph_format.line_spacing = spacing
        p.style = normal
        for run in p.runs:
            run.font.name = font; run.font.size = Pt(size_pt)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font)

    _convert_tables_to_images(doc, docx_path.stem)

    if keep_only_charts:
        removed = _filter_non_charts(doc)
        logging.info("Imágenes no gráficas eliminadas: %d", removed)

    doc.save(str(docx_path))
    logging.info("Word actualizado guardado (%s)", docx_path.name)

# ────────────────────────────────────────────────────────────────────
# Tables → PNG
# ────────────────────────────────────────────────────────────────────

def _convert_tables_to_images(doc: Document, stem: str) -> None:
    if not doc.tables:
        logging.info("No se detectaron tablas en el DOCX")
        return
    count = 0
    for tbl in list(doc.tables):
        count += 1
        data = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
        if not data:
            continue
        out_name = f"{stem}_table{count}.png"
        _df_to_png(pd.DataFrame(data[1:], columns=data[0] if len(data)>1 else None), out_name)
        # insertar imagen después de la tabla
        p = tbl._element.addnext(doc.add_paragraph()._p)  # placeholder paragraph
        pic_para = doc.paragraphs[-1]; pic_para.add_run().add_picture(out_name, width=Inches(6))
        # eliminar tabla
        tbl._element.getparent().remove(tbl._element)
        logging.info("  • %s (tabla convertida)", out_name)
    logging.info("Tablas convertidas: %d", count)


def _df_to_png(df: pd.DataFrame, out_name: str) -> None:
    fig, ax = plt.subplots(figsize=(min(8, max(2, len(df.columns)*1.2)), 0.6+0.4*len(df)))
    ax.axis('off'); tbl = ax.table(cellText=df.values, colLabels=df.columns, loc='center', cellLoc='center')
    tbl.scale(1, 1.2); fig.savefig(out_name, dpi=300, bbox_inches='tight'); plt.close(fig)

# ────────────────────────────────────────────────────────────────────
# Image heuristics
# ────────────────────────────────────────────────────────────────────

def _looks_like_chart(img: Image.Image) -> bool:
    thumb = img.resize((64, 64)).convert('RGB')
    colors = thumb.getcolors(64*64) or []
    return len(colors) < 150


def _filter_non_charts(doc: Document) -> int:
    rels = doc.part._rels; removed = 0
    for shp in list(doc.inline_shapes):
        if shp.type != WD_INLINE_SHAPE.PICTURE: continue
        img_bytes = rels[shp._inline.graphic.graphicData.pic.blipFill.blip.embed]._target._blob
        if _looks_like_chart(Image.open(BytesIO(img_bytes))):
            continue
        parent = shp._inline.getparent(); parent.getparent().remove(parent); removed += 1
    return removed

# ────────────────────────────────────────────────────────────────────
# CLI & main
# ────────────────────────────────────────────────────────────────────

def main() -> None:
    ap = argparse.ArgumentParser(add_help=False)
    ap.add_argument('pdf', nargs='?', type=Path, help='PDF de entrada')
    ap.add_argument('--include-all-images', action='store_true', help='Conservar todas las imágenes')
    ap.add_argument('--font', default='Calibri'); ap.add_argument('--size', type=float, default=11)
    ap.add_argument('--spacing', type=float, default=1.0); ap.add_argument('-h','--help', action='help')
    args = ap.parse_args()

    pdf_path = args.pdf or _pick_pdf_gui()
    if not pdf_path or not pdf_path.is_file() or pdf_path.suffix.lower() != '.pdf':
        print('[!] Selecciona un archivo .pdf válido.'); sys.exit(1)
    pdf_path = pdf_path.resolve()

    _setup_logging(pdf_path)

    try:
        word = _pdf_to_word(pdf_path)
        _extract_charts(pdf_path, args.include_all_images)
        _postprocess_docx(word, args.font, args.size, args.spacing, not args.include_all_images)
        logging.info('Proceso COMPLETADO con éxito.')
    except Exception as e:
        logging.exception('Error inesperado: %s', e)
        print('Ocurrió un error; revisa el log para detalles.')
        sys.exit(2)

    print('✓ Listo. Revisa el .docx, los .png y el log generado.')

if __name__ == '__main__':
    main()
