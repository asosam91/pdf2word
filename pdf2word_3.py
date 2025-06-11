#!/usr/bin/env python3
"""pdf_to_word_images.py
-------------------------------------------------
Convert a PDF into a Word document, apply uniform
formatting, and export embedded images as PNG files.

Outputs (for input MyFile.pdf):
    MyFile.docx              – text only, unified font/spacing, no headers/footers
    MyFile_p<page>_img<idx>.png – every image extracted at original resolution

Dependencies
------------
    pip install pdf2docx PyMuPDF python-docx

Quick usage
-----------
    python pdf_to_word_images.py MyFile.pdf \
        --font "Calibri" --size 11 --spacing 1.0

Drag‑and‑drop works on Windows (.exe build).  See README section at end
for packaging with PyInstaller.
"""
from __future__ import annotations

import argparse
from pathlib import Path
from typing import List

# PDF → DOCX
from pdf2docx import Converter  # type: ignore

# Image extraction
import fitz  # PyMuPDF

# DOCX post‑processing
from docx import Document  # type: ignore
from docx.shared import Pt  # type: ignore
from docx.oxml.ns import qn  # type: ignore


# ────────────────────────────────────────────────────────────────────
# Conversion helpers
# ────────────────────────────────────────────────────────────────────

def pdf_to_word(pdf_path: Path) -> Path:
    """Convert *pdf_path* to a DOCX stored next to it."""
    word_path = pdf_path.with_suffix(".docx")
    print(f"[+] Converting {pdf_path.name} → {word_path.name} …")
    cv = Converter(str(pdf_path))
    cv.convert(str(word_path), start=0, end=None)
    cv.close()
    return word_path


def extract_images(pdf_path: Path) -> List[str]:
    """Extract raster images from *pdf_path*; save as PNG next to the PDF."""
    doc = fitz.open(pdf_path)
    out_files: List[str] = []
    stem = pdf_path.stem

    for page_idx, page in enumerate(doc, start=1):
        for img_idx, (xref, *_rest) in enumerate(page.get_images(full=True), start=1):
            pix = fitz.Pixmap(doc, xref)
            if pix.alpha:  # flatten transparency to RGB
                pix = fitz.Pixmap(fitz.csRGB, pix)
            out_name = f"{stem}_p{page_idx}_img{img_idx}.png"
            pix.save(out_name)
            out_files.append(out_name)
            print(f"    • {out_name}")
    print(f"[+] {len(out_files)} image(s) extracted")
    return out_files


def format_docx(docx_path: Path, font: str, size_pt: float, line_spacing: float) -> None:
    """Set uniform font/size/spacing; delete headers & footers."""
    print(f"[+] Applying style ({font} {size_pt}pt, spacing {line_spacing}) …")
    doc = Document(str(docx_path))

    # Remove headers & footers
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            paragraph.clear()
        for paragraph in section.footer.paragraphs:
            paragraph.clear()

    # Normal style defines defaults
    normal_style = doc.styles["Normal"]
    normal_style.font.name = font
    normal_style.font.size = Pt(size_pt)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), font)

    # Apply to every run/paragraph
    for para in doc.paragraphs:
        para.paragraph_format.line_spacing = line_spacing
        for run in para.runs:
            run.font.name = font
            run.font.size = Pt(size_pt)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font)

    doc.save(str(docx_path))
    print("[+] Formatting done")


# ────────────────────────────────────────────────────────────────────
# CLI entry‑point
# ────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Convert PDF to Word, clean formatting, and extract images."
    )
    parser.add_argument("pdf", type=Path, help="Input PDF file path")
    parser.add_argument("--font", default="Calibri", help="Target font name (default: Calibri)")
    parser.add_argument("--size", type=float, default=11, help="Font size in points (default: 11)")
    parser.add_argument(
        "--spacing",
        type=float,
        default=1.0,
        help="Line spacing (1.0 = sencillo, 1.5 = 1.5 líneas, etc.)",
    )
    args = parser.parse_args()

    pdf_path = args.pdf.resolve()
    if not pdf_path.is_file() or pdf_path.suffix.lower() != ".pdf":
        parser.error("Input must be an existing .pdf file")

    word_file = pdf_to_word(pdf_path)
    extract_images(pdf_path)
    format_docx(word_file, args.font, args.size, args.spacing)

    print("[✓] Todo listo. Archivos guardados junto al PDF original.")
    print(f"    - {word_file.name}\n    - PNGs listados arriba")


if __name__ == "__main__":
    main()
