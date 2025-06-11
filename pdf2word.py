#!/usr/bin/env python3
"""pdf2word
============

Utilidad de línea de comandos para convertir un PDF en un documento Word con formato
uniforme y exportar las gráficas contenidas como imágenes PNG. El texto del
DOCX resultante queda con la misma fuente, tamaño, interlineado y márgenes.

Uso rápido::

    python pdf2word.py informe.pdf

Dependencias: ``pdf2docx``, ``PyMuPDF`` y ``python-docx``.
"""
from __future__ import annotations

import argparse
import logging
import sys
from io import BytesIO
from pathlib import Path
from typing import List

from pdf2docx import Converter  # type: ignore
import fitz  # PyMuPDF
from PIL import Image  # type: ignore
from docx import Document  # type: ignore
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.shape import WD_INLINE_SHAPE
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _notify(msg: str) -> None:
    """Show *msg* in a message box when bundled as executable."""
    print(msg)
    if getattr(sys, "frozen", False):  # running from PyInstaller
        try:
            import tkinter as tk
            from tkinter import messagebox

            root = tk.Tk()
            root.withdraw()
            messagebox.showinfo("pdf2word", msg)
            root.destroy()
        except Exception:
            pass

# ────────────────────────────────────────────────────────────────────────
# Logging and CLI helpers
# ────────────────────────────────────────────────────────────────────────

def _setup_logging(pdf: Path) -> None:
    log_file = pdf.with_name(f"{pdf.stem}_process.log")
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        handlers=[logging.FileHandler(log_file, encoding="utf-8"), logging.StreamHandler()],
    )
    logging.info("=== Inicio de proceso ===")


def _pick_pdf_gui() -> Path | None:
    try:
        import tkinter as tk
        from tkinter import filedialog as fd
    except Exception:
        return None
    root = tk.Tk(); root.withdraw()
    file = fd.askopenfilename(title="Selecciona un PDF", filetypes=[("PDF", "*.pdf")])
    root.destroy()
    return Path(file) if file else None

# ────────────────────────────────────────────────────────────────────────
# Conversion helpers
# ────────────────────────────────────────────────────────────────────────

def _pdf_to_word(pdf: Path) -> Path:
    """Convert *pdf* to DOCX and return the path to the created file."""
    word = pdf.with_suffix(".docx")
    logging.info("Convirtiendo %s → %s", pdf.name, word.name)
    cv = Converter(str(pdf))
    cv.convert(str(word))
    cv.close()
    return word


def _looks_like_chart(img: Image.Image) -> bool:
    """Heurística sencilla: menos de 150 colores en una miniatura."""
    thumb = img.resize((64, 64)).convert("RGB")
    colors = thumb.getcolors(64 * 64) or []
    return len(colors) < 150


def _extract_images(pdf: Path, keep_all: bool) -> List[str]:
    """Export PNG images from *pdf*.

    If ``keep_all`` is False, only images that look like charts are kept.
    Returns the list of filenames created.
    """
    doc = fitz.open(pdf)
    out: List[str] = []
    stem = pdf.stem
    for p_idx, page in enumerate(doc, 1):
        for i_idx, (xref, *_rest) in enumerate(page.get_images(full=True), 1):
            pix = fitz.Pixmap(doc, xref)
            if pix.alpha:
                pix = fitz.Pixmap(fitz.csRGB, pix)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            if not keep_all and not _looks_like_chart(img):
                continue
            label = "img" if keep_all else "chart"
            name = f"{stem}_p{p_idx}_{label}{i_idx}.png"
            pix.save(name)
            out.append(name)
            logging.info("  • %s", name)
    logging.info("Total imágenes exportadas: %d", len(out))
    return out


# ────────────────────────────────────────────────────────────────────────
# DOCX post processing
# ────────────────────────────────────────────────────────────────────────

def _clear_header_footer(part) -> None:
    """Delete paragraphs and tables in header/footer."""
    for p in list(part.paragraphs):
        p._element.getparent().remove(p._element)
    for t in list(part.tables):
        t._element.getparent().remove(t._element)


def _filter_non_charts(doc: Document) -> int:
    """Remove inline images that don't look like charts."""
    rels = doc.part._rels
    removed = 0
    for shp in list(doc.inline_shapes):
        if shp.type != WD_INLINE_SHAPE.PICTURE:
            continue
        img_bytes = rels[shp._inline.graphic.graphicData.pic.blipFill.blip.embed]._target._blob
        if _looks_like_chart(Image.open(BytesIO(img_bytes))):
            continue
        parent = shp._inline.getparent()
        parent.getparent().remove(parent)
        removed += 1
    return removed


def _apply_format(doc: Document, font: str, size: float, spacing: float, margin: float) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = font
    normal.font.size = Pt(size)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font)

    for sect in doc.sections:
        sect.top_margin = sect.bottom_margin = Inches(margin)
        sect.left_margin = sect.right_margin = Inches(margin)
        _clear_header_footer(sect.header)
        _clear_header_footer(sect.footer)

    def _format_paragraphs(paragraphs):
        for para in paragraphs:
            para.paragraph_format.line_spacing = spacing
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in para.runs:
                run.font.name = font
                run.font.size = Pt(size)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), font)

    _format_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _format_paragraphs(cell.paragraphs)


def _postprocess(docx: Path, font: str, size: float, spacing: float, margin: float, keep_charts_only: bool) -> None:
    doc = Document(str(docx))
    _apply_format(doc, font, size, spacing, margin)
    if keep_charts_only:
        removed = _filter_non_charts(doc)
        logging.info("Imágenes no gráficas eliminadas: %d", removed)
    doc.save(str(docx))
    logging.info("Word formateado y guardado (%s)", docx.name)


# ────────────────────────────────────────────────────────────────────────
# CLI
# ────────────────────────────────────────────────────────────────────────

def main() -> None:

    ap = argparse.ArgumentParser(
        description="Convertir un PDF a Word y exportar gráficas",
        add_help=False,
    )
    ap.add_argument("pdf", nargs="?", type=Path, help="Archivo PDF de entrada")
    ap.add_argument("--include-all-images", action="store_true", help="No filtrar imágenes")
    ap.add_argument("--font", default="Calibri", help="Fuente a usar")
    ap.add_argument("--size", type=float, default=11, help="Tamaño de fuente (pt)")
    ap.add_argument("--spacing", type=float, default=1.0, help="Interlineado")
    ap.add_argument("--margin", type=float, default=1.0, help="Márgenes en pulgadas")
    ap.add_argument("-h", "--help", action="help", help="Muestra esta ayuda")
    args = ap.parse_args()

    pdf = args.pdf or _pick_pdf_gui()
    if not pdf or not pdf.is_file() or pdf.suffix.lower() != ".pdf":
        _notify("[!] Selecciona un archivo .pdf válido")
        sys.exit(1)
    pdf = pdf.resolve()
    _setup_logging(pdf)

    try:
        docx = _pdf_to_word(pdf)
        _extract_images(pdf, args.include_all_images)
        _postprocess(docx, args.font, args.size, args.spacing, args.margin, not args.include_all_images)
    except Exception as exc:
        logging.exception("Error inesperado: %s", exc)
        _notify("Ocurrió un error; revisa el log para detalles")
        sys.exit(2)

    _notify("✓ Proceso completado. Revisa el DOCX, PNGs y log")


if __name__ == "__main__":
    main()
